from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer


PRIMARY = "#16324F"
MUTED = "#5D6673"
RULE = "#B8C2CC"


def parse_profile(lines: list[str]) -> dict[str, str]:
    profile = {
        "name": "Full Name",
        "email": "",
        "phone": "",
        "portfolio": "",
    }
    for line in lines:
        if line.startswith("# "):
            profile["name"] = line[2:].strip()
            continue
        for key in ["Email", "Phone", "Portfolio"]:
            prefix = f"{key}:"
            if line.startswith(prefix):
                profile[key.lower()] = line.replace(prefix, "", 1).strip()
    contact_parts = [profile["email"], profile["phone"], profile["portfolio"]]
    profile["contact_line"] = " | ".join(part for part in contact_parts if part)
    return profile


def get_subject(lines: list[str]) -> str:
    for line in lines:
        if line.startswith("Subject:"):
            return line.replace("Subject:", "").strip()
    return "Application"


def clean_line(line: str) -> str:
    return (
        line.replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .strip()
    )


def read_markdown(path: Path) -> list[str]:
    return [clean_line(line) for line in path.read_text(encoding="utf-8").splitlines()]


def add_run_text(paragraph, text: str, *, bold: bool = False, size: float | None = None):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        run = paragraph.add_run(part[2:-2] if is_bold else part)
        run.bold = bold or is_bold
        if size:
            run.font.size = Pt(size)


def build_docx(
    markdown_path: Path,
    output_path: Path,
    *,
    doc_type: str,
    profile: dict[str, str] | None = None,
):
    lines = read_markdown(markdown_path)
    profile = profile or parse_profile(lines)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, color in [
        ("Heading 1", 13.5, RGBColor(22, 50, 79)),
        ("Heading 2", 11.2, RGBColor(22, 50, 79)),
        ("Heading 3", 10.5, RGBColor(0, 0, 0)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(3)

    if doc_type == "resume":
        started = False
        role_meta_remaining = 0
        for line in lines:
            if not line:
                continue
            if line.startswith("# "):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line[2:])
                run.bold = True
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(22, 50, 79)
                continue
            if not started and not line.startswith("## "):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run_text(p, line, size=9.5)
                continue
            started = True
            if line.startswith("## "):
                role_meta_remaining = 0
                doc.add_heading(line[3:], level=1)
            elif line.startswith("### "):
                role_meta_remaining = 2
                doc.add_heading(line[4:], level=2)
            elif line.startswith("- "):
                role_meta_remaining = 0
                p = doc.add_paragraph(style="List Bullet")
                add_run_text(p, line[2:])
            else:
                p = doc.add_paragraph()
                if role_meta_remaining and len(line) <= 80:
                    run = p.add_run(line)
                    run.font.size = Pt(9.2)
                    run.font.color.rgb = RGBColor(93, 102, 115)
                    if role_meta_remaining == 2:
                        run.bold = True
                    role_meta_remaining -= 1
                else:
                    role_meta_remaining = 0
                    add_run_text(p, line)
    else:
        subject = get_subject(lines)
        body_started = False
        for line in lines:
            if not line or line.startswith("To:") or line.startswith("CC:"):
                continue
            if line.startswith("# "):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(profile["name"])
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(22, 50, 79)
                if profile["contact_line"]:
                    meta = doc.add_paragraph()
                    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_run_text(meta, profile["contact_line"], size=9)
                title = doc.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title.add_run(subject)
                title_run.bold = True
                title_run.font.size = Pt(13)
                title_run.font.color.rgb = RGBColor(22, 50, 79)
                continue
            if line.startswith("Subject:"):
                continue
            p = doc.add_paragraph()
            if not body_started:
                body_started = True
            add_run_text(p, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ResumeName",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=23,
            textColor=colors.HexColor(PRIMARY),
            alignment=1,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeCentered",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=10.5,
            textColor=colors.HexColor(MUTED),
            alignment=1,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Section",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11.6,
            leading=13.5,
            textColor=colors.HexColor(PRIMARY),
            spaceBefore=8,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Role",
            parent=styles["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=10.4,
            leading=12.2,
            textColor=colors.black,
            spaceBefore=6,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Meta",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=8.7,
            leading=10.2,
            textColor=colors.HexColor(MUTED),
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodySmall",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.1,
            leading=11.1,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletSmall",
            parent=styles["BodySmall"],
            leftIndent=14,
            firstLineIndent=-8,
            bulletIndent=0,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverName",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=22,
            textColor=colors.HexColor(PRIMARY),
            alignment=1,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverMeta",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9,
            leading=11,
            textColor=colors.HexColor(MUTED),
            alignment=1,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor(PRIMARY),
            alignment=1,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverBody",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14.5,
            spaceAfter=8,
        )
    )
    return styles


def escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("**", "")
    )


def build_pdf(
    markdown_path: Path,
    output_path: Path,
    *,
    doc_type: str,
    profile: dict[str, str] | None = None,
):
    lines = read_markdown(markdown_path)
    profile = profile or parse_profile(lines)
    styles = pdf_styles()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=0.58 * inch,
        leftMargin=0.58 * inch,
        topMargin=0.52 * inch,
        bottomMargin=0.52 * inch,
    )
    story = []

    if doc_type == "resume":
        started = False
        role_meta_remaining = 0
        for line in lines:
            if not line:
                continue
            if line.startswith("# "):
                story.append(Paragraph(escape(line[2:]), styles["ResumeName"]))
                story.append(
                    HRFlowable(
                        width="92%",
                        thickness=0.7,
                        color=colors.HexColor(RULE),
                        spaceBefore=3,
                        spaceAfter=5,
                    )
                )
            elif not started and not line.startswith("## "):
                story.append(Paragraph(escape(line), styles["ResumeCentered"]))
            else:
                started = True
                if line.startswith("## "):
                    role_meta_remaining = 0
                    story.append(Paragraph(escape(line[3:]), styles["Section"]))
                    story.append(
                        HRFlowable(
                            width="100%",
                            thickness=0.45,
                            color=colors.HexColor(RULE),
                            spaceBefore=0,
                            spaceAfter=3,
                        )
                    )
                elif line.startswith("### "):
                    role_meta_remaining = 2
                    story.append(Paragraph(escape(line[4:]), styles["Role"]))
                elif line.startswith("- "):
                    role_meta_remaining = 0
                    story.append(Paragraph(escape(line[2:]), styles["BulletSmall"], bulletText="-"))
                else:
                    if role_meta_remaining and len(line) <= 80:
                        story.append(Paragraph(escape(line), styles["Meta"]))
                        role_meta_remaining -= 1
                    else:
                        role_meta_remaining = 0
                        story.append(Paragraph(escape(line), styles["BodySmall"]))
    else:
        subject = get_subject(lines)
        for line in lines:
            if line.startswith("Subject:"):
                continue
            if not line or line.startswith("To:") or line.startswith("CC:"):
                continue
            if line.startswith("# "):
                story.append(Paragraph(escape(profile["name"]), styles["CoverName"]))
                if profile["contact_line"]:
                    story.append(Paragraph(escape(profile["contact_line"]), styles["CoverMeta"]))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=0.6,
                        color=colors.HexColor(RULE),
                        spaceBefore=0,
                        spaceAfter=12,
                    )
                )
                story.append(Paragraph(escape(subject), styles["CoverTitle"]))
                story.append(Spacer(1, 0.05 * inch))
            else:
                story.append(Paragraph(escape(line), styles["CoverBody"]))

    doc.build(story)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--resume-md", required=True)
    parser.add_argument("--cover-md", required=True)
    parser.add_argument("--out-dir", required=True)
    parser.add_argument("--prefix", required=True)
    args = parser.parse_args()

    resume_md = Path(args.resume_md)
    cover_md = Path(args.cover_md)
    out_dir = Path(args.out_dir)
    prefix = args.prefix
    profile = parse_profile(read_markdown(resume_md))

    build_docx(resume_md, out_dir / f"{prefix} CV.docx", doc_type="resume", profile=profile)
    build_pdf(resume_md, out_dir / f"{prefix} CV.pdf", doc_type="resume", profile=profile)
    build_docx(cover_md, out_dir / f"{prefix} Cover Letter.docx", doc_type="cover", profile=profile)
    build_pdf(cover_md, out_dir / f"{prefix} Cover Letter.pdf", doc_type="cover", profile=profile)


if __name__ == "__main__":
    main()
